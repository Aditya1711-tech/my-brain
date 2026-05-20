import io

from app.services.documents.types import RawExtraction


def parse_docx(file_bytes: bytes, filename: str) -> RawExtraction:
    """Extract text from a DOCX file."""
    import docx

    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)

    # Extract tables
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            tables.append(rows)

    return RawExtraction(
        text=text,
        tables=tables,
        page_count=1,  # docx doesn't have page concept easily
    )
